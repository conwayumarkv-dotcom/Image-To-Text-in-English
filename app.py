import streamlit as st
from google import genai
from google.genai import types
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import io
import time
import threading

# ==========================================
# 1. 워드 파일 디자인 & 생성 헬퍼 함수
# ==========================================
def set_cell_borders(cell, color="D9D9D9", sz="4", val="single"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
    tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def create_word_document(all_word_data):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("통합 본문 단어장")
    title_run.font.bold = True
    title_run.font.size = Pt(18)
    title_p.paragraph_format.space_after = Pt(24)
    
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    col_widths = [Inches(1.5), Inches(2.0), Inches(4.0)]
    
    hdr_cells = table.rows[0].cells
    headers = ["본문 단어", "우리말 뜻", "영영 풀이"]
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        hdr_cells[i].width = col_widths[i]
        set_cell_shading(hdr_cells[i], "4F81BD")
        set_cell_borders(hdr_cells[i], color="A6A6A6")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = docx.shared.RGBColor(255, 255, 255)
            
    for row_idx, item in enumerate(all_word_data):
        row_cells = table.add_row().cells
        row_data = [item.get("word", ""), item.get("meaning", ""), item.get("definition", "")]
        
        for i, text in enumerate(row_data):
            # [TypeError 예외 차단 완료] text 데이터가 없거나 형식이 숫자인 경우를 대비해 안전하게 문자열 변환 처리
            row_cells[i].text = str(text) if text is not None else ""
            row_cells[i].width = col_widths[i]
            if row_idx % 2 == 1:
                set_cell_shading(row_cells[i], "F2F5F8")
            set_cell_borders(row_cells[i], color="D9D9D9")
            
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(5)
            for run in p.runs:
                run.font.size = Pt(10)
                
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ==========================================
# 2. 백그라운드 AI 호출 및 에러 자동 재시도 워커 함수
# ==========================================
def gemini_api_worker(client, image_bytes, mime_type, prompt, result_container):
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[
                    types.Part.from_bytes(data=image_bytes, mime_type=mime_type),
                    prompt
                ],
                config=types.GenerateContentConfig(response_mime_type="application/json")
            )
            result_container["data"] = json.loads(response.text)
            result_container["status"] = "success"
            return
        except Exception as e:
            error_msg = str(e)
            if ("503" in error_msg or "unavailable" in error_msg.lower() or "demand" in error_msg.lower()) and attempt < max_retries - 1:
                time.sleep(2.0 * (attempt + 1))
                continue
            
            result_container["status"] = "error"
            result_container["error_msg"] = error_msg
            return

# ==========================================
# 3. 이미지 비동기 분석 및 균일 선형 정밀 제어 로직
# ==========================================
def process_images_safely(client, uploaded_files, api_key, progress_bar, status_text):
    all_data = []
    total_files = len(uploaded_files)
    
    prompt = """
    이 이미지에서 영어 단어, 우리말 뜻, 영영 풀이를 추출해서 정확한 JSON 배열 형식으로 출력해줘.
    필기구로 수정한 흔적이나 추가로 적은 필기는 무시하고, 원래 인쇄되어 있던 텍스트만 추출해줘.
    결과는 오직 아래 구조를 가진 JSON 데이터만 반환해야 해:
    [
      {"word": "단어", "meaning": "품사 및 뜻", "definition": "영영 풀이 내용"}
    ]
    """
    
    ui_progress = 0.0
    
    for idx, file in enumerate(uploaded_files):
        file.seek(0)
        image_bytes = file.read()
        
        worker_result = {"status": "pending", "data": None, "error_msg": None}
        
        api_thread = threading.Thread(
            target=gemini_api_worker, 
            args=(client, image_bytes, file.type, prompt, worker_result)
        )
        api_thread.start()
        
        start_progress = idx / total_files
        target_max_progress = (idx + 1) / total_files
        file_share = 1.0 / total_files
        
        # [정체 현상 해결 핵심 알고리즘]
        # AI 응답 대기 시간 동안 끊김을 방지하기 위해 파일당 정해진 진행 속도를 아주 정밀하게 쪼개어 배정합니다.
        # 루프당 고정 증가폭을 균일화하여 50% 구간이나 대기 마감 지점(95% 이상)에서의 멈춤을 완벽 제거합니다.
        step_increment = file_share * 0.006 
        
        while api_thread.is_alive():
            # 다음 파일 도달 목표치의 97% 지점까지는 흔들림 없는 일정한 대각선 선형 속도로 부드럽게 상승시킵니다.
            if ui_progress < (start_progress + (file_share * 0.97)):
                ui_progress += step_increment
            else:
                # 혹시 AI 분석 속도가 가상 타임라인보다 살짝 늦어질 경우에만 초미세 속도로 계속 전진하도록 유도합니다.
                ui_progress += file_share * 0.0005
                
            if ui_progress > 0.99: ui_progress = 0.99
            
            progress_bar.progress(ui_progress)
            status_text.markdown(f"🔍 **[ {int(ui_progress * 100)}% / 100% ]** ({idx+1}/{total_files}장째) AI가 사진 속 영어 단어를 열심히 읽어내고 있어요..")
            time.sleep(0.05)
            
        if worker_result["status"] == "success" and worker_result["data"]:
            all_data.extend(worker_result["data"])
            
            # AI 데이터 처리가 응답된 직후, 가쁜 정체 없이 다음 파일 파트로 바톤을 정밀하게 연결해 넘겨줍니다.
            while ui_progress < target_max_progress:
                ui_progress += 0.02
                if ui_progress > target_max_progress: ui_progress = target_max_progress
                progress_bar.progress(ui_progress)
                status_text.markdown(f"✨ **[ {int(ui_progress * 100)}% / 100% ]** ({idx+1}/{total_files}장째) 선생님 단어장에 맞게 예쁘게 다듬는 중입니다!")
                time.sleep(0.01)
                
        elif worker_result["status"] == "error":
            error_msg = worker_result["error_msg"]
            if "quota" in error_msg.lower() or "limit" in error_msg.lower():
                if idx > 0:
                    st.warning("⚠️ 오늘 준비된 무료 변환량(20장)을 모두 사용하셨습니다. 아쉽지만 현재까지 성공한 단어들로만 Word 문서를 만듭니다.")
                    break
                else:
                    st.error("❌ 구글이 제공하는 무료 하루 사용량(20장)을 초과하여 지금은 변환할 수 없습니다. 내일 다시 이용해 주세요.")
                    return None
            elif "503" in error_msg or "unavailable" in error_msg.lower():
                st.error("❌ 순간적으로 구글 AI 서버에 사용자가 몰려 응답이 지연되었습니다. 잠시 후 'Word 파일로 변환하기' 버튼을 한 번만 더 눌러주세요.")
                return None
            else:
                st.error(f"❌ 단어 변환 중 예상치 못한 오류가 생겼습니다: {error_msg}")
                return None
                
    if all_data:
        progress_bar.progress(1.0)
        status_text.success("🌿 **[ 100% / 100% ]** 수업용 단어장이 완성되었습니다! 아래 다운로드 버튼을 눌러보세요!")
    return all_data

# ==========================================
# 4. Streamlit 메인 UI 대시보드
# ==========================================
st.set_page_config(page_title="Voca-converter", layout="centered", page_icon="📝")

st.markdown("""
    <style>
    html, body, [data-testid="stAppViewContainer"] {
        background-color: #FBF9F4 !important;
    }
    
    [data-testid="stMainBlockContainer"] {
        background-color: transparent !important;
        max-width: 720px !important;
        margin: 0 auto !important;
        padding-top: 50px !important;
    }
    
    [data-testid="stVerticalBlockBorderContainer"] {
        border: none !important;
        background: transparent !important;
        box-shadow: none !important;
        padding: 0 !important;
    }

    .brand-title {
        font-size: 52px !important;
        font-weight: 700 !important;
        color: #556B2F !important;
        text-align: center !important;
        margin-bottom: 5px !important;
        letter-spacing: -1px !important;
    }
    
    .brand-caption {
        font-size: 15px !important;
        color: #8C9A86 !important;
        text-align: center !important;
        margin-bottom: 5px !important;
        font-weight: 500 !important;
    }
    
    .brand-author {
        font-size: 11px !important;
        color: #A0ABA2 !important;
        text-align: right !important;
        margin-bottom: 45px !important;
        font-weight: 500 !important;
        padding-right: 5px;
        letter-spacing: 0.5px;
    }

    [data-testid="stFileUploader"] {
        border: none !important;
        background-color: #EEF1F6 !important;
        border-radius: 14px !important;
        padding: 20px 25px !important;
    }
    
    div.stButton > button:first-child {
        background-color: #85A392 !important; 
        color: white !important;
        border: none !important;
        font-size: 16px !important;
        font-weight: 500 !important;
        border-radius: 10px !important;
        padding: 12px 24px !important;
        box-shadow: none !important;
        width: auto !important;
    }
    div.stButton > button:first-child:hover {
        background-color: #6C8B7A !important;
    }
    
    [data-testid="stDownloadButton"]>button {
        background-color: #78909C !important;
        color: white !important;
        border-radius: 10px !important;
        border: none !important;
        padding: 12px 24px !important;
    }
    [data-testid="stDownloadButton"]>button:hover {
        background-color: #607D8B !important;
    }
    
    div[data-testid="stNotification"] {
        background-color: #E8F1FC !important;
        border: none !important;
        border-radius: 12px !important;
    }
    div[data-testid="stNotification"] p {
        color: #1E60B4 !important;
        font-weight: 500 !important;
    }
    
    .stProgress > div > div > div > div {
        background-color: #85A392 !important;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown("<div class='brand-title'>Voca-converter</div>", unsafe_allow_html=True)
st.markdown("<div class='brand-caption'>사진 속 지문을 인식하여 편집 가능한 워드 문서(.docx)로 변환합니다.</div>", unsafe_allow_html=True)
st.markdown("<div class='brand-author'>© TOP English Academy. All rights reserved.</div>", unsafe_allow_html=True)

if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    st.error("❌ Streamlit Cloud 설정의 Secrets에 GEMINI_API_KEY가 등록되지 않았습니다.")
    st.stop()

uploaded_files = st.file_uploader(
    "변환할 영어 지문 사진을 업로드하세요 (복수 선택 가능)", 
    type=["jpg", "jpeg", "png"], 
    accept_multiple_files=True
)

if uploaded_files:
    st.write("")
    st.markdown(f"📂 **{len(uploaded_files)}개의 파일이 선택되었습니다.**")
    
    if st.button("Word 파일로 변환하기 ✨", type="primary"):
        client = genai.Client(api_key=api_key)
        
        status_text = st.empty()
        progress_bar = st.progress(0)
        
        all_word_data = process_images_safely(client, uploaded_files, api_key, progress_bar, status_text)
        
        if all_word_data:
            st.toast("단어 데이터 정제가 완료되었습니다!")
            st.write("---")
            st.write("### 🔍 데이터 통합 미리보기")
            st.dataframe(all_word_data, use_container_width=True)
            
            word_file_buffer = create_word_document(all_word_data)
            
            st.download_button(
                label="📥 수업용 영어 단어장 워드파일(.docx) 다운로드 받기",
                data=word_file_buffer,
                file_name="🔮_통합_영어단어장.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
